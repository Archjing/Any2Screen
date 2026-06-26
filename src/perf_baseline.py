from __future__ import annotations

import json
import os
import platform
import subprocess
import sys
import tempfile
import time
from dataclasses import asdict, dataclass
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from weasyprint import HTML


@dataclass(frozen=True)
class BaselineScenario:
    scenario_id: str
    title: str
    input_kind: str
    output_format: str
    stage_category: str
    input_path: Path
    commands: tuple[tuple[str, ...], ...]
    expected_output: Path
    notes: str


@dataclass(frozen=True)
class EnvironmentSummary:
    name: str
    recorded_at: str
    machine: str
    python_version: str
    runtime: str
    notes: str


@dataclass(frozen=True)
class ScenarioResult:
    scenario_id: str
    title: str
    stage_category: str
    input_kind: str
    output_format: str
    command: str
    success: bool
    duration_seconds: float
    output_bytes: int | None
    output_width: int | None
    output_height: int | None
    output_pages: int | None
    output_path: str
    error_message: str


def ensure_sample_inputs(sample_root: Path) -> dict[str, Path]:
    inputs_dir = sample_root / "inputs"
    inputs_dir.mkdir(parents=True, exist_ok=True)

    markdown_path = inputs_dir / "article.md"
    if not markdown_path.exists():
        markdown_path.write_text(
            "# Any2Screen Baseline Sample\n\n"
            "这是一份用于性能基线的 Markdown 样例。\n\n"
            "## Summary\n\n"
            "- mobile reading\n"
            "- pdf export\n"
            "- image export\n\n"
            "```python\nprint('baseline')\n```\n",
            encoding="utf-8",
        )

    text_path = inputs_dir / "notes.txt"
    if not text_path.exists():
        text_path.write_text(
            "Any2Screen baseline text sample.\n"
            "Line two keeps the TXT -> HTML conversion path stable.\n",
            encoding="utf-8",
        )

    docx_path = inputs_dir / "brief.docx"
    if not docx_path.exists():
        document = Document()
        document.add_heading("Any2Screen DOCX Baseline", level=1)
        document.add_paragraph("This DOCX fixture feeds the doc2md + any2html path.")
        document.save(docx_path)

    pdf_path = inputs_dir / "brief.pdf"
    if not pdf_path.exists():
        HTML(
            string=(
                "<html><body><h1>Any2Screen PDF Baseline</h1>"
                "<p>This PDF fixture feeds the doc2md + any2html path.</p>"
                "</body></html>"
            )
        ).write_pdf(str(pdf_path))

    return {
        "markdown": markdown_path,
        "text": text_path,
        "docx": docx_path,
        "pdf": pdf_path,
    }


def build_default_scenarios(sample_root: Path) -> list[BaselineScenario]:
    inputs = ensure_sample_inputs(sample_root)
    return [
        BaselineScenario(
            scenario_id="markdown-html",
            title="Markdown -> HTML",
            input_kind="markdown",
            output_format="html",
            stage_category="html_generation",
            input_path=inputs["markdown"],
            commands=(("scripts/any2screen.py", "any2html", str(inputs["markdown"]), "-o", "out"),),
            expected_output=Path("out/article.html"),
            notes="Direct any2html stage using fixed Markdown sample.",
        ),
        BaselineScenario(
            scenario_id="text-html",
            title="TXT -> HTML",
            input_kind="text",
            output_format="html",
            stage_category="html_generation",
            input_path=inputs["text"],
            commands=(
                ("scripts/any2screen.py", "doc2md", str(inputs["text"]), "-o", "stage"),
                ("scripts/any2screen.py", "any2html", "stage/notes.md", "-o", "out"),
            ),
            expected_output=Path("out/notes.html"),
            notes="Measures doc2md preprocessing plus HTML generation for TXT.",
        ),
        BaselineScenario(
            scenario_id="docx-html",
            title="DOCX -> HTML",
            input_kind="docx",
            output_format="html",
            stage_category="html_generation",
            input_path=inputs["docx"],
            commands=(
                ("scripts/any2screen.py", "doc2md", str(inputs["docx"]), "-o", "stage"),
                ("scripts/any2screen.py", "any2html", "stage/brief.md", "-o", "out"),
            ),
            expected_output=Path("out/brief.html"),
            notes="Measures doc2md preprocessing plus HTML generation for DOCX.",
        ),
        BaselineScenario(
            scenario_id="pdf-html",
            title="PDF -> HTML",
            input_kind="pdf",
            output_format="html",
            stage_category="html_generation",
            input_path=inputs["pdf"],
            commands=(
                ("scripts/any2screen.py", "doc2md", str(inputs["pdf"]), "-o", "stage"),
                ("scripts/any2screen.py", "any2html", "stage/brief.md", "-o", "out"),
            ),
            expected_output=Path("out/brief.html"),
            notes="Measures doc2md preprocessing plus HTML generation for PDF.",
        ),
        BaselineScenario(
            scenario_id="markdown-pdf-a4",
            title="Markdown -> PDF (A4)",
            input_kind="markdown",
            output_format="pdf",
            stage_category="pdf_render",
            input_path=inputs["markdown"],
            commands=(("scripts/any2screen.py", "convert", str(inputs["markdown"]), "--pdf", "-o", "out"),),
            expected_output=Path("out/article.pdf"),
            notes="Standard A4 PDF render through the main pipeline.",
        ),
        BaselineScenario(
            scenario_id="markdown-pdf-wechat",
            title="Markdown -> WeChat PDF",
            input_kind="markdown",
            output_format="wechat-pdf",
            stage_category="pdf_render",
            input_path=inputs["markdown"],
            commands=(("scripts/any2screen.py", "convert", str(inputs["markdown"]), "--wechat", "-o", "out"),),
            expected_output=Path("out/article.wechat.pdf"),
            notes="Narrow mobile PDF render through the main pipeline.",
        ),
        BaselineScenario(
            scenario_id="markdown-image-small-png",
            title="Markdown -> PNG (small screen)",
            input_kind="markdown",
            output_format="png",
            stage_category="image_render",
            input_path=inputs["markdown"],
            commands=(
                (
                    "scripts/any2screen.py",
                    "convert",
                    str(inputs["markdown"]),
                    "--img",
                    "--width",
                    "430",
                    "--format",
                    "png",
                    "-o",
                    "out",
                ),
            ),
            expected_output=Path("out/article.png"),
            notes="Small-screen long image render for mobile sharing.",
        ),
        BaselineScenario(
            scenario_id="markdown-image-large-jpeg",
            title="Markdown -> JPEG (large screen)",
            input_kind="markdown",
            output_format="jpeg",
            stage_category="image_render",
            input_path=inputs["markdown"],
            commands=(
                (
                    "scripts/any2screen.py",
                    "convert",
                    str(inputs["markdown"]),
                    "--img",
                    "--width",
                    "1080",
                    "--format",
                    "jpeg",
                    "-o",
                    "out",
                ),
            ),
            expected_output=Path("out/article.jpeg"),
            notes="Large-screen long image render for higher-fidelity sharing.",
        ),
    ]


def summarize_local_environment(name: str, runtime: str, notes: str) -> EnvironmentSummary:
    return EnvironmentSummary(
        name=name,
        recorded_at=datetime.now().astimezone().isoformat(timespec="seconds"),
        machine=platform.platform(),
        python_version=platform.python_version(),
        runtime=runtime,
        notes=notes,
    )


def run_scenario(repo_root: Path, scenario: BaselineScenario, python_executable: str = sys.executable) -> ScenarioResult:
    with tempfile.TemporaryDirectory(prefix=f"any2screen-{scenario.scenario_id}-") as tmp:
        workdir = Path(tmp)
        command_strings: list[str] = []
        last_output = ""
        started = time.perf_counter()
        try:
            for command in scenario.commands:
                resolved_command = list(command)
                resolved_command[0] = str(repo_root / resolved_command[0])
                argv = [python_executable, *resolved_command]
                command_strings.append(" ".join(argv))
                completed = subprocess.run(
                    argv,
                    cwd=workdir,
                    capture_output=True,
                    text=True,
                    env={**os.environ, "PYTHONPATH": str(repo_root / "src")},
                )
                last_output = "\n".join(part for part in (completed.stdout, completed.stderr) if part).strip()
                if completed.returncode != 0:
                    error = last_output
                    raise RuntimeError(error or f"command failed with exit code {completed.returncode}")
            duration = time.perf_counter() - started
            output_path = workdir / scenario.expected_output
            if not output_path.exists():
                detail = f"expected output not found: {output_path}"
                if last_output:
                    detail = f"{detail}; command output: {last_output}"
                raise FileNotFoundError(detail)
            output_bytes = output_path.stat().st_size
            width, height = read_dimensions(output_path)
            pages = read_pdf_page_count(output_path)
            return ScenarioResult(
                scenario_id=scenario.scenario_id,
                title=scenario.title,
                stage_category=scenario.stage_category,
                input_kind=scenario.input_kind,
                output_format=scenario.output_format,
                command=" && ".join(command_strings),
                success=True,
                duration_seconds=round(duration, 3),
                output_bytes=output_bytes,
                output_width=width,
                output_height=height,
                output_pages=pages,
                output_path=str(scenario.expected_output),
                error_message="",
            )
        except Exception as exc:
            duration = time.perf_counter() - started
            return ScenarioResult(
                scenario_id=scenario.scenario_id,
                title=scenario.title,
                stage_category=scenario.stage_category,
                input_kind=scenario.input_kind,
                output_format=scenario.output_format,
                command=" && ".join(command_strings) if command_strings else "",
                success=False,
                duration_seconds=round(duration, 3),
                output_bytes=None,
                output_width=None,
                output_height=None,
                output_pages=None,
                output_path=str(scenario.expected_output),
                error_message=str(exc),
            )


def read_dimensions(path: Path) -> tuple[int | None, int | None]:
    if path.suffix.lower() == ".png":
        with path.open("rb") as handle:
            data = handle.read(24)
        if data[:8] != b"\x89PNG\r\n\x1a\n":
            return None, None
        return int.from_bytes(data[16:20], "big"), int.from_bytes(data[20:24], "big")
    if path.suffix.lower() in {".jpg", ".jpeg"}:
        with path.open("rb") as handle:
            data = handle.read()
        index = 2
        while index < len(data):
            if data[index] != 0xFF:
                index += 1
                continue
            marker = data[index + 1]
            index += 2
            if marker in {0xC0, 0xC1, 0xC2, 0xC3}:
                block_length = int.from_bytes(data[index:index + 2], "big")
                height = int.from_bytes(data[index + 3:index + 5], "big")
                width = int.from_bytes(data[index + 5:index + 7], "big")
                return width, height
            block_length = int.from_bytes(data[index:index + 2], "big")
            index += block_length
        return None, None
    return None, None


def read_pdf_page_count(path: Path) -> int | None:
    if path.suffix.lower() != ".pdf":
        return None
    try:
        from PyPDF2 import PdfReader

        return len(PdfReader(str(path)).pages)
    except Exception:
        return None


def render_markdown_report(
    environments: Iterable[EnvironmentSummary],
    grouped_results: dict[str, list[ScenarioResult]],
    docker_command: str,
) -> str:
    lines = [
        "# Any2Screen Performance Baseline",
        "",
        "## Scope",
        "",
        "- Baseline covers Markdown, TXT, DOCX, PDF input samples with fixed output parameters.",
        "- Timing is grouped into `html_generation`, `pdf_render`, and `image_render`.",
        "- D27 re-test must compare the same scenario IDs, command parameters, output byte size, page/image dimensions, and success state.",
        "- Regression rule: any scenario slower by more than `20%` or failing after a previous success counts as a regression unless documented with an approved reason.",
        "",
        "## Environment",
        "",
    ]

    for environment in environments:
        lines.extend(
            [
                f"### {environment.name}",
                "",
                f"- Recorded at: `{environment.recorded_at}`",
                f"- Machine: `{environment.machine}`",
                f"- Python: `{environment.python_version}`",
                f"- Runtime: `{environment.runtime}`",
                f"- Notes: {environment.notes}",
                "",
            ]
        )

    for environment_name, results in grouped_results.items():
        lines.extend(
            [
                f"## Results: {environment_name}",
                "",
                "| Scenario | Stage | Input | Output | Success | Duration (s) | Bytes | Dimensions | Pages | Output |",
                "| --- | --- | --- | --- | --- | ---: | ---: | --- | ---: | --- |",
            ]
        )
        for result in results:
            dimensions = "-"
            if result.output_width and result.output_height:
                dimensions = f"{result.output_width}x{result.output_height}"
            pages = "-" if result.output_pages is None else str(result.output_pages)
            size = "-" if result.output_bytes is None else str(result.output_bytes)
            lines.append(
                "| {scenario} | {stage} | {input_kind} | {output} | {success} | {duration:.3f} | {bytes_} | {dimensions} | {pages} | `{path}` |".format(
                    scenario=result.scenario_id,
                    stage=result.stage_category,
                    input_kind=result.input_kind,
                    output=result.output_format,
                    success="yes" if result.success else "no",
                    duration=result.duration_seconds,
                    bytes_=size,
                    dimensions=dimensions,
                    pages=pages,
                    path=result.output_path,
                )
            )
            lines.append(f"  Command: `{result.command}`")
            if result.error_message:
                lines.append(f"  Error: `{result.error_message}`")
        lines.append("")

    lines.extend(
        [
            "## Docker Collection",
            "",
            "Use the following command to collect the Docker baseline with the same scenario set:",
            "",
            f"`{docker_command}`",
            "",
            "If Docker collection cannot run on the current machine, record the concrete blocker and date instead of marking the baseline complete silently.",
            "",
            "## D27 Comparison Fields",
            "",
            "- `scenario_id`",
            "- `stage_category`",
            "- command parameters",
            "- `duration_seconds`",
            "- `output_bytes`",
            "- `output_pages`",
            "- `output_width` / `output_height`",
            "- `success` and `error_message`",
        ]
    )
    return "\n".join(lines) + "\n"


def dump_json(path: Path, environments: Iterable[EnvironmentSummary], grouped_results: dict[str, list[ScenarioResult]]) -> None:
    payload = {
        "environments": [asdict(environment) for environment in environments],
        "results": {name: [asdict(result) for result in results] for name, results in grouped_results.items()},
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
