import unittest
from tempfile import TemporaryDirectory
from unittest.mock import patch

from tests import _paths  # noqa: F401

try:
    import fastapi  # noqa: F401
except ImportError:
    HAS_FASTAPI = False
else:
    HAS_FASTAPI = True


@unittest.skipIf(not HAS_FASTAPI, "FastAPI dependencies are not installed")
class WebApiTests(unittest.TestCase):
    def test_create_app_registers_api_routes(self) -> None:
        from web import create_app

        app = create_app()
        paths = {route.path for route in app.routes}

        self.assertIn("/", paths)
        self.assertIn("/api/exports/{file_id}/html", paths)
        self.assertIn("/api/exports/{file_id}/image", paths)
        self.assertIn("/api/exports/{file_id}/pdf", paths)
        self.assertIn("/api/exports/{file_id}/wechat-pdf", paths)
        self.assertIn("/api/files", paths)
        self.assertIn("/api/health", paths)
        self.assertIn("/api/previews/{file_id}", paths)
        self.assertIn("/api/previews/{file_id}/html", paths)
        self.assertIn("/api/version", paths)
        self.assertEqual(app.title, "Any2Screen API")

    def test_create_app_mounts_static_assets(self) -> None:
        from web import create_app

        app = create_app()
        names = {route.name for route in app.routes}

        self.assertIn("assets", names)

    def test_health_response_shape(self) -> None:
        from web.routes import health

        payload = health().model_dump()

        self.assertEqual(payload["status"], "ok")
        self.assertIn("timestamp", payload)

    def test_version_response_shape(self) -> None:
        from web.routes import version

        payload = version().model_dump()

        self.assertEqual(payload["name"], "any2screen")
        self.assertEqual(payload["api_version"], "0.1.0")
        self.assertIn("upload", payload["capabilities"])
        self.assertIn("preview", payload["capabilities"])

    def test_file_registry_detects_supported_markdown_upload(self) -> None:
        from web.files import FileRegistry

        with TemporaryDirectory() as tmp:
            registry = FileRegistry(upload_root=tmp)
            record = registry.add("report.md", b"# Report")

            self.assertEqual(record.filename, "report.md")
            self.assertEqual(record.size_bytes, 8)
            self.assertEqual(record.extension, ".md")
            self.assertEqual(record.detected_type, "markdown")
            self.assertTrue(record.supported)
            self.assertEqual(record.content, b"# Report")
            self.assertEqual(record.path.read_bytes(), b"# Report")
            self.assertEqual(record.path.parent.parent.as_posix(), tmp)
            self.assertIs(registry.get(record.file_id), record)

    def test_file_registry_marks_unknown_uploads(self) -> None:
        from web.files import FileRegistry

        registry = FileRegistry()
        record = registry.add("archive.zip", b"PK")

        self.assertEqual(record.detected_type, "unknown")
        self.assertFalse(record.supported)

    def test_preview_file_generates_markdown_preview(self) -> None:
        from web.files import file_registry
        from web.routes import preview_file

        record = file_registry.add("report.md", b"# Report\n\nVisible\n\nHidden")
        payload = preview_file(record.file_id, blocks=2).model_dump()

        self.assertEqual(payload["file_id"], record.file_id)
        self.assertEqual(payload["detected_type"], "markdown")
        self.assertTrue(payload["truncated"])
        self.assertIn("<title>Report</title>", payload["html"])
        self.assertIn("Visible", payload["html"])
        self.assertNotIn("Hidden", payload["html"])

    def test_preview_file_generates_text_preview(self) -> None:
        from web.files import file_registry
        from web.routes import preview_file

        record = file_registry.add("notes.txt", b"plain text")
        payload = preview_file(record.file_id).model_dump()

        self.assertEqual(payload["detected_type"], "text")
        self.assertIn("<title>notes</title>", payload["html"])
        self.assertIn("plain text", payload["html"])

    def test_preview_file_generates_docx_preview(self) -> None:
        from io import BytesIO

        from docx import Document
        from web.files import file_registry
        from web.routes import preview_file

        buffer = BytesIO()
        document = Document()
        document.add_heading("Docx Title", level=1)
        document.add_paragraph("Docx paragraph")
        document.save(buffer)

        record = file_registry.add("sample.docx", buffer.getvalue())
        payload = preview_file(record.file_id).model_dump()

        self.assertEqual(payload["detected_type"], "docx")
        self.assertIn("<title>Docx Title</title>", payload["html"])
        self.assertIn("Docx paragraph", payload["html"])

    def test_preview_file_generates_pdf_preview(self) -> None:
        from weasyprint import HTML
        from web.files import file_registry
        from web.routes import preview_file

        content = HTML(string="<h1>PDF Title</h1><p>PDF paragraph</p>").write_pdf()
        record = file_registry.add("sample.pdf", content)
        payload = preview_file(record.file_id).model_dump()

        self.assertEqual(payload["detected_type"], "pdf")
        self.assertIn("<title>sample</title>", payload["html"])
        self.assertIn("PDF", payload["html"])

    def test_preview_html_returns_html_response(self) -> None:
        from web.files import file_registry
        from web.routes import preview_html

        record = file_registry.add("report.md", b"# Report")
        response = preview_html(record.file_id)

        self.assertEqual(response.media_type, "text/html; charset=utf-8")
        self.assertIn(b'<base href="/">', response.body)
        self.assertIn(b"<title>Report</title>", response.body)

    def test_export_html_returns_download_response(self) -> None:
        from web.files import file_registry
        from web.routes import export_html_file

        record = file_registry.add("report.md", b"# Report")
        response = export_html_file(record.file_id)
        output_path = record.path.with_suffix(".html")

        self.assertEqual(response.media_type, "text/html; charset=utf-8")
        self.assertIn('filename="report.html"', response.headers["Content-Disposition"])
        self.assertIn(b"<title>Report</title>", response.body)
        self.assertEqual(output_path.parent, record.path.parent)
        self.assertTrue(output_path.exists())
        self.assertIn("<title>Report</title>", output_path.read_text(encoding="utf-8"))

    def test_export_pdf_returns_download_response(self) -> None:
        from web.files import file_registry
        from web.routes import export_pdf_file

        record = file_registry.add("report.md", b"# Report")
        response = export_pdf_file(record.file_id)
        output_path = record.path.with_suffix(".pdf")

        self.assertEqual(response.media_type, "application/pdf")
        self.assertIn('filename="report.pdf"', response.headers["Content-Disposition"])
        self.assertTrue(response.body.startswith(b"%PDF"))
        self.assertEqual(output_path.parent, record.path.parent)
        self.assertTrue(output_path.exists())
        self.assertTrue(output_path.read_bytes().startswith(b"%PDF"))

    def test_export_wechat_pdf_returns_download_response(self) -> None:
        from web.files import file_registry
        from web.routes import export_wechat_pdf_file

        record = file_registry.add("report.md", b"# Report")
        response = export_wechat_pdf_file(record.file_id)
        output_path = record.path.with_suffix(".wechat.pdf")

        self.assertEqual(response.media_type, "application/pdf")
        self.assertIn('filename="report.wechat.pdf"', response.headers["Content-Disposition"])
        self.assertTrue(response.body.startswith(b"%PDF"))
        self.assertEqual(output_path.parent, record.path.parent)
        self.assertTrue(output_path.exists())
        self.assertTrue(output_path.read_bytes().startswith(b"%PDF"))

    def test_export_image_returns_download_response(self) -> None:
        from web.files import file_registry
        from web.routes import export_image_file

        def fake_render_image(html_path, image_path, width, image_format, verbose=False):
            image_path.write_bytes(b"\x89PNG\r\n\x1a\nfake")
            return True, image_path.stat().st_size, width, 1200

        record = file_registry.add("report.md", b"# Report")
        with patch("web.export.render_image", side_effect=fake_render_image) as render:
            response = export_image_file(record.file_id, screen="small", format="png")
        output_path = record.path.with_name("report.small.png")

        self.assertEqual(response.media_type, "image/png")
        self.assertIn('filename="report.small.png"', response.headers["Content-Disposition"])
        self.assertTrue(response.body.startswith(b"\x89PNG"))
        self.assertEqual(output_path.parent, record.path.parent)
        self.assertTrue(output_path.exists())
        render.assert_called_once()
        self.assertEqual(render.call_args.kwargs["width"], 430)
        self.assertEqual(render.call_args.kwargs["image_format"], "png")

    def test_export_image_large_jpeg_uses_large_screen_preset(self) -> None:
        from web.files import file_registry
        from web.routes import export_image_file

        def fake_render_image(html_path, image_path, width, image_format, verbose=False):
            image_path.write_bytes(b"\xff\xd8fake")
            return True, image_path.stat().st_size, width, 1600

        record = file_registry.add("report.md", b"# Report")
        with patch("web.export.render_image", side_effect=fake_render_image) as render:
            response = export_image_file(record.file_id, screen="large", format="jpeg")
        output_path = record.path.with_name("report.large.jpeg")

        self.assertEqual(response.media_type, "image/jpeg")
        self.assertIn('filename="report.large.jpeg"', response.headers["Content-Disposition"])
        self.assertTrue(response.body.startswith(b"\xff\xd8"))
        self.assertTrue(output_path.exists())
        render.assert_called_once()
        self.assertEqual(render.call_args.kwargs["width"], 1080)
        self.assertEqual(render.call_args.kwargs["image_format"], "jpeg")

    def test_export_image_rejects_invalid_query_params(self) -> None:
        try:
            from fastapi.testclient import TestClient
        except RuntimeError as e:
            self.skipTest(f"FastAPI TestClient dependencies are not installed: {e}")
        from web import create_app
        from web.files import file_registry

        record = file_registry.add("report.md", b"# Report")
        client = TestClient(create_app())

        self.assertEqual(client.get(f"/api/exports/{record.file_id}/image?screen=tiny").status_code, 422)
        self.assertEqual(client.get(f"/api/exports/{record.file_id}/image?format=webp").status_code, 422)

    def test_export_rejects_unsupported_type(self) -> None:
        from fastapi import HTTPException
        from web.files import file_registry
        from web.routes import export_html_file

        record = file_registry.add("image.png", b"fake")

        with self.assertRaises(HTTPException) as ctx:
            export_html_file(record.file_id)

        self.assertEqual(ctx.exception.status_code, 415)

    def test_preview_file_rejects_unsupported_type(self) -> None:
        from fastapi import HTTPException
        from web.files import file_registry
        from web.routes import preview_file

        record = file_registry.add("image.png", b"fake")

        with self.assertRaises(HTTPException) as ctx:
            preview_file(record.file_id)

        self.assertEqual(ctx.exception.status_code, 415)


if __name__ == "__main__":
    unittest.main()
